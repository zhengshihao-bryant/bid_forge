# -*- coding: utf-8 -*-
"""
generation/assembler.py —— M4-09 完整标书组装

    BidDocumentAssembler.assemble(tender_id) → {markdown, docx_path, ...}

前序组装（章节树平铺）→ Markdown（封面/目录/章节正文/响应表/元数据）
→ DOCX（python-docx：封面 + 页眉页脚 + 段落/表格 + 中文字体 eastAsia 宋体）。

- CH-02 目录：按章节树前序自动生成（占位页码）。
- CH-08 需求响应表：用 M4-07 三列响应表（覆盖全部规范需求）替换占位模板。
- 未生成章节：渲染占位文案，不阻断整单（M4-10 断点继续的产物可重新组装）。
"""
from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import Optional

from ... import config
from ...db import Database
from ...schemas import now_str
from .job import GenerationJobRunner
from .models import SectionStatus
from .outline import OutlineBuilder, tree_from_flat
from .response_table import BidResponseTableBuilder

logger = logging.getLogger(__name__)

DEFAULT_OUTPUT_DIR = config.DATA_DIR / "out"


class BidDocumentAssembler:
    """M4-09 文档组装器（Markdown + DOCX 双产物）。"""

    def __init__(self, db: Optional[Database] = None):
        self.db = db or Database(config.DB_PATH)

    # ------------------------------------------------------------------
    def assemble(self, tender_id: str,
                 output_dir: Optional[Path] = None) -> dict:
        """前序组装 → Markdown + DOCX。返回 {markdown, docx_path, 统计}。"""
        sections = self._ordered_sections(tender_id)
        if not sections:
            raise ValueError("该招标项目尚未规划章节（请先 POST /outline）")
        md = self.render_markdown(tender_id, sections)
        docx_path = self.render_docx(tender_id, sections, output_dir)
        done = sum(1 for s in sections
                   if s.status in (SectionStatus.DONE, SectionStatus.SKIPPED))
        return {
            "tender_id": tender_id,
            "total_sections": len(sections),
            "done_sections": done,
            "markdown": md,
            "docx_path": str(docx_path),
            "markdown_chars": len(md),
        }

    # ------------------------------------------------------------------
    # Markdown
    # ------------------------------------------------------------------
    def render_markdown(self, tender_id: str, sections: Optional[list] = None,
                        ) -> str:
        sections = sections or self._ordered_sections(tender_id)
        parts: list[str] = []
        for sec in sections:
            parts.append(self._section_md(tender_id, sec))
        parts.append(self._metadata_md(tender_id, sections))
        return "\n\n---\n\n".join(parts)

    def _section_md(self, tender_id: str, sec) -> str:
        if sec.id == "CH-02":
            return self._render_toc(sec)
        if sec.id == "CH-08":
            return BidResponseTableBuilder(self.db).to_markdown(tender_id)
        row = self.db.query_one("SELECT * FROM generation_sections "
                                "WHERE section_id = ?", (sec.id,))
        if row and row.get("content_md"):
            return row["content_md"]
        return f"## {sec.title}\n\n（本章节未生成，请先运行生成任务）"

    def _render_toc(self, sec) -> str:
        lines = [f"## {sec.title}", ""]
        for s in self._ordered_sections(sec.tender_id):
            if s.id == "CH-02":
                continue
            indent = "  " * (s.level - 1)
            lines.append(f"{indent}{s.ord}. {s.title}")
        return "\n".join(lines)

    def _metadata_md(self, tender_id: str, sections: list) -> str:
        done = sum(1 for s in sections
                   if s.status in (SectionStatus.DONE, SectionStatus.SKIPPED))
        failed = sum(1 for s in sections if s.status == SectionStatus.FAILED)
        warns = 0
        rows = self.db.query(
            "SELECT warnings FROM generation_sections WHERE tender_id = ?",
            (tender_id,))
        for r in rows:
            warns += len(json.loads(r.get("warnings") or "[]"))
        tender = self.db.query_one("SELECT name FROM tenders WHERE id = ?",
                                   (tender_id,))
        return (f"## 生成信息\n\n"
                f"- 招标项目：{tender['name'] if tender else tender_id}\n"
                f"- 章节：{done}/{len(sections)} 完成，{failed} 失败\n"
                f"- 事实校验告警：{warns} 条（详见各章节，供 M5 一致性核查）\n"
                f"- 生成时间：{now_str()}")

    # ------------------------------------------------------------------
    # DOCX（python-docx 1.2.0 程序化组装）
    # ------------------------------------------------------------------
    def render_docx(self, tender_id: str, sections: Optional[list] = None,
                    output_dir: Optional[Path] = None) -> Path:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt

        sections = sections or self._ordered_sections(tender_id)
        out_dir = output_dir or DEFAULT_OUTPUT_DIR
        out_dir.mkdir(parents=True, exist_ok=True)
        path = out_dir / f"{tender_id}_投标文件.docx"

        doc = Document()
        _set_normal_font(doc)
        self._add_cover(doc, tender_id)
        self._add_header_footer(doc, tender_id, qn, OxmlElement)

        for sec in sections:
            if sec.id in ("CH-01", "CH-02"):
                continue                       # 封面 / 目录已单独渲染
            if sec.id == "CH-08":
                self._add_response_table_docx(doc, tender_id)
                continue
            draft = self._load_draft(sec.id)
            if draft is None or not draft.paragraphs:
                doc.add_heading(sec.title, level=1)
                doc.add_paragraph("（本章节未生成，请先运行生成任务）")
                continue
            self._add_draft_docx(doc, draft)

        self._add_metadata_docx(doc, tender_id)
        doc.save(path)
        logger.info("DOCX 已生成: %s（%d 字节）", path, path.stat().st_size)
        return path

    def _load_draft(self, section_id: str):
        row = self.db.query_one("SELECT * FROM generation_sections "
                                "WHERE section_id = ?", (section_id,))
        if not row or not (row.get("content_md") or row.get("paragraphs")):
            return None
        return Database.row_to_section(row)

    def _add_cover(self, doc, tender_id: str) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        tender = self.db.query_one("SELECT name FROM tenders WHERE id = ?",
                                   (tender_id,))
        name = tender["name"] if tender else tender_id
        for _ in range(4):
            doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(name)
        r.font.size = Pt(26); r.bold = True
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("投标文件")
        r.font.size = Pt(22); r.bold = True
        doc.add_paragraph()
        for line in ("招标编号：＿＿＿＿＿＿", "投标人：＿＿＿＿＿＿＿＿",
                     "日期：＿＿＿＿年＿＿月＿＿日"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(line)
        doc.add_page_break()

    def _add_header_footer(self, doc, tender_id: str, qn, OxmlElement) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        tender = self.db.query_one("SELECT name FROM tenders WHERE id = ?",
                                   (tender_id,))
        name = tender["name"] if tender else tender_id
        header = doc.sections[0].header
        hp = header.paragraphs[0]
        hp.text = f"{name} —— 投标文件"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in hp.runs:
            _set_run_font(run, qn)
        footer = doc.sections[0].footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run("第 ")
        _set_run_font(run, qn)
        run2 = fp.add_run("1")
        _add_page_field(run2, qn, OxmlElement)
        _set_run_font(run2, qn)
        run3 = fp.add_run(" 页")
        _set_run_font(run3, qn)

    def _add_draft_docx(self, doc, draft) -> None:
        doc.add_heading(draft.title, level=1)
        for p in draft.paragraphs:
            if p.type == "heading":
                doc.add_heading(p.text, level=min(max(p.level or 2, 2), 4))
            elif p.type == "list_item":
                doc.add_paragraph(p.text, style="List Bullet")
            elif p.type == "table":
                self._add_table_docx(doc, p.table)
            else:
                doc.add_paragraph(p.text)
        if draft.warnings:
            doc.add_paragraph("")
            w = doc.add_paragraph("⚠ 事实校验告警（供 M5 一致性核查）：")
            w.bold = True
            for msg in draft.warnings:
                doc.add_paragraph(f"- {msg}")

    def _add_table_docx(self, doc, rows: list) -> None:
        if not rows:
            return
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = "Table Grid"
        for i, row in enumerate(rows):
            for j, cell in enumerate(row):
                para = table.cell(i, j).paragraphs[0]
                run = para.add_run(str(cell))
                if i == 0:
                    run.bold = True

    def _add_response_table_docx(self, doc, tender_id: str) -> None:
        built = BidResponseTableBuilder(self.db).build(tender_id)
        doc.add_heading("需求响应表", level=1)
        doc.add_paragraph(
            f"共 {built['total']} 条规范需求："
            + "；".join(f"{k} {v}" for k, v in built["counts"].items() if v))
        rows = [["招标要求", "企业响应", "证据"]]
        for r in built["rows"][:50]:
            req = f"{r['title']}（{r['req_type']}）\n{r['text']}"
            rows.append([req, r["response"], r["evidence_ids"] and
                         r["evidences"][0]["evidence_id"] or "—"])
        self._add_table_docx(doc, rows)
        doc.add_paragraph("MISSING/UNKNOWN 不编造：响应列如实陈述，具体数值以【待确认】标注。")

    def _add_metadata_docx(self, doc, tender_id: str) -> None:
        doc.add_page_break()
        doc.add_heading("生成信息", level=1)
        assembled = self.assemble_meta(tender_id)
        for k, v in assembled.items():
            doc.add_paragraph(f"{k}：{v}")

    def assemble_meta(self, tender_id: str) -> dict:
        sections = self._ordered_sections(tender_id)
        done = sum(1 for s in sections
                   if s.status in (SectionStatus.DONE, SectionStatus.SKIPPED))
        failed = sum(1 for s in sections if s.status == SectionStatus.FAILED)
        return {"章节": f"{done}/{len(sections)} 完成，{failed} 失败",
                "生成时间": now_str()}

    # ------------------------------------------------------------------
    def _ordered_sections(self, tender_id: str) -> list:
        rows = self.db.query(
            "SELECT * FROM generation_sections WHERE tender_id = ?",
            (tender_id,))
        flat = [Database.row_to_bid_section(r) for r in rows]
        return OutlineBuilder.flatten(tree_from_flat(flat))


# ---------------------------------------------------------------------------
# DOCX 字体/页码工具
# ---------------------------------------------------------------------------
def _set_normal_font(doc) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    for sname in ("Normal", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        try:
            style = doc.styles[sname]
        except KeyError:
            continue
        style.font.name = "宋体"
        if sname == "Normal":
            style.font.size = Pt(12)
        # 中文字体 eastAsia（无此设置，Word 中文标题可能回退字体/乱码）
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), "宋体")


def _set_run_font(run, qn) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _add_page_field(run, qn, OxmlElement) -> None:
    """脚注页码域：第 {PAGE} 页（静态目录占位页码，域在 Word 中刷新）。"""
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


__all__ = ["BidDocumentAssembler", "DEFAULT_OUTPUT_DIR"]
