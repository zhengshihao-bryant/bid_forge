# -*- coding: utf-8 -*-
"""
tests/test_m4_docx.py —— M4-09 文档组装（批次 4）

覆盖（对照 M4-11）：
- 章节顺序组装：前序（公司概况先于技术方案；封面在最前）
- Markdown：封面/目录/章节正文/三列响应表/生成信息
- DOCX：文件存在、可被 python-docx 打开、段落>0、含资质表、中文字体 eastAsia
- 未生成章节 → 占位文案不阻断
- API：GET /document（markdown + docx FileResponse）
"""
from __future__ import annotations

import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "backend"))

from app.services.generation import (BidDocumentAssembler,  # noqa: E402
                                     GenerationJobRunner)


@pytest.fixture()
def generated(seed_m4, tmp_path):
    """seed_m4 + 全量生成 + 组装输出到 tmp_path。"""
    data = seed_m4
    runner = GenerationJobRunner(db=data["db"])
    job = runner.create_job(data["tender_id"])
    runner.run(job)
    assert job.status == "已完成"
    return data, tmp_path


def _markdown_of(data, tmp_path):
    return BidDocumentAssembler(db=data["db"]).assemble(
        data["tender_id"], output_dir=tmp_path)["markdown"]


# ═══════════════════════════════════════════════════════════════════════
# Markdown 组装
# ═══════════════════════════════════════════════════════════════════════
def test_markdown_preorder_section_order(generated):
    """前序组装：封面→目录→商务→技术→实施→售后→响应表。"""
    data, tmp_path = generated
    md = _markdown_of(data, tmp_path)
    assert md.startswith("## 封面") or "## 封面" in md
    idx = {
        "封面": md.index("## 封面"),
        "目录": md.index("## 目录"),
        "公司概况": md.index("## 公司概况与综合实力"),
        "技术方案": md.index("## 总体技术方案"),
        "实施": md.index("## 项目实施计划"),
        "售后": md.index("## 售后服务承诺"),
        "响应表": md.index("# 需求响应表"),      # CH-08 用单井号标题（M4-07 产物）
    }
    assert idx["封面"] < idx["目录"] < idx["公司概况"] < idx["技术方案"] \
        < idx["实施"] < idx["售后"] < idx["响应表"]


def test_markdown_contains_core_content(generated):
    data, tmp_path = generated
    md = _markdown_of(data, tmp_path)
    # 公司概况数值、资质编号、人员、质保事实（能力卡回填的确定性内容）
    assert "注册资本5000万元" in md
    assert "ISO9001" in md and "等保三级" in md
    assert "张伟" in md
    assert "质保期3年" in md
    # 响应表三列 + 状态口径
    assert "| 招标要求 | 企业响应 | 证据 |" in md
    assert "MISSING=资料明确显示不满足" in md
    # 生成信息（元数据）
    assert "## 生成信息" in md
    assert "26" in md


def test_markdown_placeholder_for_ungenerated(seed_m4, tmp_path):
    """未生成章节 → 占位文案，不阻断组装。"""
    data = seed_m4
    md = BidDocumentAssembler(db=data["db"]).assemble(
        data["tender_id"], output_dir=tmp_path)["markdown"]
    assert "（本章节未生成，请先运行生成任务）" in md


# ═══════════════════════════════════════════════════════════════════════
# DOCX 组装
# ═══════════════════════════════════════════════════════════════════════
def test_docx_file_exists_and_opens(generated):
    """DOCX：文件存在、可打开、段落>0、含表格、中文内容完整。"""
    from docx import Document

    data, tmp_path = generated
    result = BidDocumentAssembler(db=data["db"]).assemble(
        data["tender_id"], output_dir=tmp_path)
    path = Path(result["docx_path"])
    assert path.exists() and path.stat().st_size > 0
    doc = Document(str(path))
    texts = [p.text for p in doc.paragraphs]
    assert len(texts) > 10, "文档段落过少"
    joined = "\n".join(texts)
    # 封面 + 章节正文 + 生成信息
    assert "投标文件" in joined
    assert "注册资本5000万元" in joined
    assert "张伟" in joined
    assert "生成信息" in joined
    # 表格（资质表 / 指标响应表 / 响应表至少一张）
    assert len(doc.tables) >= 3, f"应有资质/指标/响应表，实际 {len(doc.tables)}"
    # 表格含资质编号
    flat = " ".join(cell.text for t in doc.tables for row in t.rows
                    for cell in row.cells)
    assert "ISO9001" in flat and "等保三级" in flat


def test_docx_east_asia_font(generated):
    """中文字体：Normal 样式含 w:eastAsia=宋体（防乱码）。"""
    from docx.oxml.ns import qn

    from app.services.generation.assembler import _set_normal_font
    from docx import Document

    data, tmp_path = generated
    doc = Document()
    _set_normal_font(doc)
    rPr = doc.styles["Normal"].element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    assert rFonts.get(qn("w:eastAsia")) == "宋体"


def test_docx_tables_render_rows(generated):
    """响应表 DOCX：三列表头 + 需求行。"""
    from docx import Document

    data, tmp_path = generated
    result = BidDocumentAssembler(db=data["db"]).assemble(
        data["tender_id"], output_dir=tmp_path)
    doc = Document(str(result["docx_path"]))
    resp_tables = [t for t in doc.tables
                   if [c.text for c in t.rows[0].cells]
                   == ["招标要求", "企业响应", "证据"]]
    assert resp_tables, "应有三列响应表"
    assert len(resp_tables[0].rows) >= 30, "响应表应覆盖规范需求行"


# ═══════════════════════════════════════════════════════════════════════
# API
# ═══════════════════════════════════════════════════════════════════════
@pytest.fixture()
def doc_client(seed_m4, tmp_path):
    from fastapi.testclient import TestClient
    from app.api.main import app

    data = seed_m4
    data["db"].insert("tenders", {"id": data["tender_id"], "name": "M4文档测试项目",
                                  "created_at": "2026-01-01 00:00:00"})
    runner = GenerationJobRunner(db=data["db"])
    job = runner.create_job(data["tender_id"])
    runner.run(job)
    with TestClient(app) as c:
        yield data, c


def test_document_endpoint_markdown(doc_client):
    data, c = doc_client
    r = c.get(f"/api/generation/tenders/{data['tender_id']}/document")
    assert r.status_code == 200
    j = r.json()
    assert j["total_sections"] == 26
    assert "## 封面" in j["markdown"]
    assert "| 招标要求 | 企业响应 | 证据 |" in j["markdown"]


def test_document_endpoint_docx_file(doc_client):
    data, c = doc_client
    r = c.get(f"/api/generation/tenders/{data['tender_id']}/document"
              "?format=docx")
    assert r.status_code == 200
    assert r.headers["content-type"].startswith(
        "application/vnd.openxmlformats")
    assert len(r.content) > 0
