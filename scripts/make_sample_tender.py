# -*- coding: utf-8 -*-
"""
scripts/make_sample_tender.py —— 《XX市智慧园区建设项目》样例招标文件包生成器

产出（backend/data/samples/智慧园区项目/）：

  01_招标文件正文.docx      约 110 页：13 章（LLM 生成正文 + 规则生成评分表/格式章）
  02_技术规格书.pdf         约 30 页：PyMuPDF 排版（simsun.ttc/simhei.ttf + 目录书签）
  03_设备清单.xlsx          3 工作表约 150 行：openpyxl 规则生成
  04_补充通知(扫描件).pdf    2 页：PIL 渲染文字成图 → 无文本层 PDF（OCR 测试靶）
  样例说明.md               文件清单 + 预埋需求基线（M1 验收对照）

用法：

  python scripts/make_sample_tender.py                 # 完整模式（LLM 生成正文，可续跑）
  python scripts/make_sample_tender.py --no-llm        # 离线模式（内置简版文本，不调 API）
  python scripts/make_sample_tender.py --only pdf      # 只生成指定类别：docx/pdf/xlsx/scan/readme
  python scripts/make_sample_tender.py --no-cache      # 忽略缓存重新生成

设计：
- 预埋需求基线：样例内容中按清单种入量化要求（人脸识别/设备接入≥1000/项目经理5年/评分50+20+30…），
  供 M1 需求提取验收对照（见 样例说明.md）
- 生成缓存 scripts/_sample_cache.json（gitignored）：重复运行不重复计费、断点续跑
- 扫描件无文本层：pdf_parser 应检测出 ocr_pages=[1,2]
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

BACKEND = Path(__file__).resolve().parent.parent / "backend"
sys.path.insert(0, str(BACKEND))

from app import config  # noqa: E402  (加载 .env + 控制台 UTF-8 兜底)
from app.services.llm import create_llm_client  # noqa: E402

OUT_DIR = config.SAMPLES_DIR / "智慧园区项目"
CACHE_FILE = Path(__file__).resolve().parent / "_sample_cache.json"

TENDER_NAME = "XX市智慧园区建设项目"
TENDER_NO = "ZH-2024-018"

# ═══════════════════════════════════════════════════════════════════════
# 正文章节定义：(章节标题, [生成单元])；单元 = (键, 小节标题, 要点清单, 目标字数)
# 要点清单即预埋需求种入点 —— LLM 提示词要求写全，样例说明.md 与其对照
# ═══════════════════════════════════════════════════════════════════════
CHAPTER_UNITS: list[tuple[str, list[tuple[str, str, list[str], int]]]] = [
    ("第二章 投标人须知", [
        ("2-1", "2.1 总则与项目概况", [
            "招标人：XX市智慧城市管理局；项目名称：XX市智慧园区建设项目；招标编号：ZH-2024-018",
            "项目资金来源为财政资金，已落实；招标方式为公开招标",
            "招标范围：园区智能化软硬件一体化建设（含安装、调试、培训、质保期服务）",
        ], 2200),
        ("2-2", "2.2 招标文件的组成、澄清与修改", [
            "招标文件组成：招标公告、投标人须知、项目背景与建设目标、技术要求、功能要求、实施要求、人员要求、资质要求、售后服务要求、商务要求、评分标准、投标文件格式、合同主要条款",
            "投标人对招标文件有疑问的，应在投标截止时间 15 日前以书面形式提出澄清",
            "招标人对已发出的招标文件进行必要的澄清或者修改的，应在投标截止时间 15 日前以书面形式通知所有获取招标文件的潜在投标人",
        ], 2200),
        ("2-3", "2.3 投标文件的编制与递交", [
            "投标文件应使用中文编制，胶装成册；正本 1 份、副本 4 份，另附电子版 1 份（U 盘）",
            "投标截止时间：2024 年 10 月 15 日 9 时 30 分；开标时间与地点：同投标截止时间，XX市公共资源交易中心第二开标室",
            "投标有效期：自投标截止之日起 90 天；投标保证金：人民币 20 万元",
        ], 2200),
    ]),
    ("第三章 项目背景与建设目标", [
        ("3-1", "3.1 园区现状与建设背景", [
            "园区占地约 800 亩，入驻企业 120 余家，园区内工作人员约 1.5 万人",
            "现状问题：安防、门禁、停车、能耗等系统各自为政、数据孤岛、管理效率低",
            "建设背景：落实城市数字化转型要求，打造智慧园区示范标杆",
        ], 2200),
        ("3-2", "3.2 建设目标与建设原则", [
            "建设目标：实现园区管理'一屏统览、一网统管、一键联动'",
            "建成覆盖安防、通行、能耗、停车、一卡通的综合管理平台",
            "建设原则：统一规划、分步实施、安全可控、适度超前、开放兼容",
        ], 2000),
    ]),
    ("第四章 技术要求", [
        ("4-1", "4.1 总体架构要求", [
            "采用'感知层-网络层-平台层-应用层'分层架构",
            "平台层采用微服务架构，支持横向扩展；建设统一数据中台",
            "对外提供标准 REST API 与消息接口，接口开放可集成",
        ], 2200),
        ("4-2", "4.2 平台功能要求", [
            "平台应支持人脸识别，实现园区人员身份识别及统一管理（★条款）",
            "平台应支持不少于 1000 台（个）设备的接入管理",
            "平台应支持不少于 500 个并发用户的同时在线访问",
            "平台应支持统一身份认证与分级权限管理",
        ], 2400),
        ("4-3", "4.3 数据与安全要求", [
            "数据要求：统一数据标准、数据共享交换、重要数据每日备份",
            "安全要求：系统安全保护等级不低于网络安全等级保护第三级（等保三级）",
            "应满足商用密码应用安全性评估要求；提供完整审计日志",
        ], 2200),
        ("4-4", "4.4 性能、接口与国产化要求", [
            "系统可用性不低于 99.9%；一般业务操作响应时间不超过 3 秒",
            "视频监控接口应符合 GB/T 28181 标准；设备接入支持 ONVIF",
            "系统应支持信创国产化环境部署：国产 CPU、国产操作系统、国产数据库、国产中间件",
        ], 2200),
    ]),
    ("第五章 功能要求", [
        ("5-1", "5.1 智慧安防与智慧通行系统", [
            "视频监控：400 万像素以上网络摄像机，视频存储不少于 90 天",
            "门禁管理：支持刷卡、人脸、二维码等多种认证方式",
            "人脸识别通行：支持白名单比对与陌生人预警；访客管理支持线上预约",
        ], 2400),
        ("5-2", "5.2 智慧能耗与智慧停车系统", [
            "能耗管理：水、电分类分项计量与在线监测，支持能耗分析与告警",
            "智能照明：公共区域照明分时分区控制",
            "停车管理：车牌识别无感通行、车位引导、无人值守收费",
        ], 2200),
        ("5-3", "5.3 一卡通系统与综合管理平台", [
            "一卡通：门禁、消费、考勤、访客一卡通用",
            "综合管理平台：三维可视化呈现，一屏统览园区运行态势",
            "提供移动端应用（APP/小程序）",
        ], 2200),
    ]),
    ("第六章 实施要求", [
        ("6-1", "6.1 工期与实施组织要求", [
            "总工期不超过 12 个月（自合同签订之日起计算）",
            "中标人应成立专门项目管理机构，项目经理须常驻现场",
            "实施前须提交详细实施组织方案与进度计划，经招标人批准后执行",
        ], 2000),
        ("6-2", "6.2 进度、培训与验收要求", [
            "进度计划应设置关键里程碑节点，按月报送进度报告",
            "培训要求：系统管理员、操作人员、运维人员分类培训，总学时不少于 40 学时",
            "验收：初验合格后试运行 3 个月，试运行无重大问题后组织终验",
        ], 2000),
    ]),
    ("第七章 人员要求", [
        ("7-1", "7.1 项目团队与人员要求", [
            "项目团队配置：项目经理、技术负责人、实施工程师、培训讲师等",
            "★项目经理须具有 5 年以上智慧园区类项目管理经验，并具有 PMP 或信息系统项目管理师证书",
            "技术负责人须具有高级工程师职称；项目实施期间驻场人员不少于 5 人",
        ], 1800),
    ]),
    ("第八章 资质要求", [
        ("8-1", "8.1 投标人资格、资质与业绩要求", [
            "投标人须为独立法人，注册资本不低于 1000 万元人民币",
            "★投标人须通过 ISO9001 质量管理体系认证、ISO27001 信息安全管理体系认证，并具有 CMMI3 级及以上认证",
            "业绩要求：近三年（2021 年 1 月以来）至少完成 3 个类似智慧园区或智慧城市项目，单个合同金额不低于 500 万元",
            "投标人无失信记录，未被列入政府采购严重违法失信行为记录名单",
        ], 2000),
    ]),
    ("第九章 售后服务要求", [
        ("9-1", "9.1 售后服务要求", [
            "质保期不少于 2 年（自终验合格之日起计算）",
            "故障响应时间不超过 2 小时到达现场；质保期内免费维修、免费更换故障部件",
            "质保期内提供 2 名驻场运维人员；免费提供备品备件清单并常备易损件",
        ], 1800),
    ]),
    ("第十章 商务要求", [
        ("10-1", "10.1 商务与报价要求", [
            "报价方式：总价包干（含设备、软件、安装调试、培训、税费等一切费用）",
            "投标人须提交分项报价表；投标保证金 2%；履约保证金为合同总价的 5%",
            "付款方式：预付款 30%、验收合格后支付 65%、质保金 5% 于质保期满后无息退还",
        ], 2000),
    ]),
    ("第十三章 合同主要条款", [
        ("13-1", "13.1 合同主要条款", [
            "合同书格式：采用招标文件提供的合同书格式，双方协商一致可另行签订补充协议",
            "知识产权：本项目产生的软件著作权及技术成果归招标人所有",
            "违约责任：逾期交付违约金按每日合同总价的 0.05% 计算；逾期超过 30 天招标人有权解除合同",
        ], 2000),
    ]),
]

# 离线模式内置简版文本（预埋需求同样种入，保证 --no-llm 下管线可验收）
FALLBACK_TEXT: dict[str, str] = {
    "2-1": "### 2.1 总则与项目概况\n招标人：XX市智慧城市管理局。项目名称：XX市智慧园区建设项目，招标编号：ZH-2024-018。项目资金来源为财政资金，已落实；招标方式为公开招标。招标范围：园区智能化软硬件一体化建设（含安装、调试、培训、质保期服务）。",
    "2-2": "### 2.2 招标文件的组成、澄清与修改\n招标文件组成：招标公告、投标人须知、项目背景与建设目标、技术要求、功能要求、实施要求、人员要求、资质要求、售后服务要求、商务要求、评分标准、投标文件格式、合同主要条款。投标人对招标文件有疑问的，应在投标截止时间 15 日前以书面形式提出澄清。招标人对已发出的招标文件进行必要的澄清或者修改的，应在投标截止时间 15 日前以书面形式通知所有获取招标文件的潜在投标人。",
    "2-3": "### 2.3 投标文件的编制与递交\n投标文件应使用中文编制，胶装成册；正本 1 份、副本 4 份，另附电子版 1 份（U 盘）。投标截止时间：2024 年 10 月 15 日 9 时 30 分；开标时间与地点：同投标截止时间，XX市公共资源交易中心第二开标室。投标有效期：自投标截止之日起 90 天。投标保证金：人民币 20 万元。",
    "3-1": "### 3.1 园区现状与建设背景\n园区占地约 800 亩，入驻企业 120 余家，园区内工作人员约 1.5 万人。现状问题：安防、门禁、停车、能耗等系统各自为政、数据孤岛、管理效率低。建设背景：落实城市数字化转型要求，打造智慧园区示范标杆。",
    "3-2": "### 3.2 建设目标与建设原则\n建设目标：实现园区管理'一屏统览、一网统管、一键联动'，建成覆盖安防、通行、能耗、停车、一卡通的综合管理平台。建设原则：统一规划、分步实施、安全可控、适度超前、开放兼容。",
    "4-1": "### 4.1 总体架构要求\n系统采用'感知层-网络层-平台层-应用层'分层架构。平台层采用微服务架构，支持横向扩展；建设统一数据中台。对外提供标准 REST API 与消息接口，接口开放可集成。",
    "4-2": "### 4.2 平台功能要求\n（1）平台应支持人脸识别，实现园区人员身份识别及统一管理（★条款）。（2）平台应支持不少于 1000 台（个）设备的接入管理。（3）平台应支持不少于 500 个并发用户的同时在线访问。（4）平台应支持统一身份认证与分级权限管理。",
    "4-3": "### 4.3 数据与安全要求\n数据要求：统一数据标准、数据共享交换、重要数据每日备份。安全要求：系统安全保护等级不低于网络安全等级保护第三级（等保三级）。应满足商用密码应用安全性评估要求，提供完整审计日志。",
    "4-4": "### 4.4 性能、接口与国产化要求\n系统可用性不低于 99.9%；一般业务操作响应时间不超过 3 秒。视频监控接口应符合 GB/T 28181 标准；设备接入支持 ONVIF。系统应支持信创国产化环境部署：国产 CPU、国产操作系统、国产数据库、国产中间件。",
    "5-1": "### 5.1 智慧安防与智慧通行系统\n视频监控：400 万像素以上网络摄像机，视频存储不少于 90 天。门禁管理：支持刷卡、人脸、二维码等多种认证方式。人脸识别通行：支持白名单比对与陌生人预警；访客管理支持线上预约。",
    "5-2": "### 5.2 智慧能耗与智慧停车系统\n能耗管理：水、电分类分项计量与在线监测，支持能耗分析与告警。智能照明：公共区域照明分时分区控制。停车管理：车牌识别无感通行、车位引导、无人值守收费。",
    "5-3": "### 5.3 一卡通系统与综合管理平台\n一卡通：门禁、消费、考勤、访客一卡通用。综合管理平台：三维可视化呈现，一屏统览园区运行态势。提供移动端应用（APP/小程序）。",
    "6-1": "### 6.1 工期与实施组织要求\n总工期不超过 12 个月（自合同签订之日起计算）。中标人应成立专门项目管理机构，项目经理须常驻现场。实施前须提交详细实施组织方案与进度计划，经招标人批准后执行。",
    "6-2": "### 6.2 进度、培训与验收要求\n进度计划应设置关键里程碑节点，按月报送进度报告。培训要求：系统管理员、操作人员、运维人员分类培训，总学时不少于 40 学时。验收：初验合格后试运行 3 个月，试运行无重大问题后组织终验。",
    "7-1": "### 7.1 项目团队与人员要求\n项目团队配置：项目经理、技术负责人、实施工程师、培训讲师等。★项目经理须具有 5 年以上智慧园区类项目管理经验，并具有 PMP 或信息系统项目管理师证书。技术负责人须具有高级工程师职称；项目实施期间驻场人员不少于 5 人。",
    "8-1": "### 8.1 投标人资格、资质与业绩要求\n投标人须为独立法人，注册资本不低于 1000 万元人民币。★投标人须通过 ISO9001 质量管理体系认证、ISO27001 信息安全管理体系认证，并具有 CMMI3 级及以上认证。业绩要求：近三年（2021 年 1 月以来）至少完成 3 个类似智慧园区或智慧城市项目，单个合同金额不低于 500 万元。投标人无失信记录，未被列入政府采购严重违法失信行为记录名单。",
    "9-1": "### 9.1 售后服务要求\n质保期不少于 2 年（自终验合格之日起计算）。故障响应时间不超过 2 小时到达现场；质保期内免费维修、免费更换故障部件。质保期内提供 2 名驻场运维人员；免费提供备品备件清单并常备易损件。",
    "10-1": "### 10.1 商务与报价要求\n报价方式：总价包干（含设备、软件、安装调试、培训、税费等一切费用）。投标人须提交分项报价表；投标保证金 2%；履约保证金为合同总价的 5%。付款方式：预付款 30%、验收合格后支付 65%、质保金 5% 于质保期满后无息退还。",
    "13-1": "### 13.1 合同主要条款\n合同书格式：采用招标文件提供的合同书格式，双方协商一致可另行签订补充协议。知识产权：本项目产生的软件著作权及技术成果归招标人所有。违约责任：逾期交付违约金按每日合同总价的 0.05% 计算；逾期超过 30 天招标人有权解除合同。",
}


# ═══════════════════════════════════════════════════════════════════════
# 缓存
# ═══════════════════════════════════════════════════════════════════════
def load_cache() -> dict:
    if CACHE_FILE.exists():
        return json.loads(CACHE_FILE.read_text(encoding="utf-8"))
    return {}


def save_cache(cache: dict) -> None:
    CACHE_FILE.write_text(json.dumps(cache, ensure_ascii=False, indent=2), encoding="utf-8")


# ═══════════════════════════════════════════════════════════════════════
# 生成单元内容（LLM / 缓存 / 离线降级）
# ═══════════════════════════════════════════════════════════════════════
def generate_unit(client, key: str, section: str, points: list[str],
                  chars: int, use_llm: bool, cache: dict) -> str:
    if key in cache and cache[key].get("text"):
        print(f"  [缓存] {section}")
        return cache[key]["text"]
    if not use_llm:
        print(f"  [离线] {section}")
        return FALLBACK_TEXT.get(key, f"### {section}\n（离线模式简版内容）")

    point_lines = "\n".join(f"- {p}" for p in points)
    user = (
        f"你是智慧园区建设项目的招标文件编制人员（甲方口吻）。\n"
        f"请撰写「{section}」一节，使用正式、规范的招标文件语言。\n"
        f"必须完整包含以下要点（量化数字原样保留，不得改动）：\n{point_lines}\n"
        f"输出 markdown 文本：以 \"### {section}\" 开头，正文用自然段（每段 80-150 字），"
        f"要点可以合并到段落中，总长度约 {chars} 字。只输出正文，不要输出任何解释或前言。"
    )
    system = ("你是专业的政府采购招标文件编制人员，熟悉智慧园区/智慧城市项目。"
              "输出必须是 markdown 正文（json 无关，纯文本），语言规范、数字严谨。")
    print(f"  [LLM] {section} ...")
    resp = client.chat([
        {"role": "system", "content": system},
        {"role": "user", "content": user},
    ], max_tokens=3000)
    text = (resp.get("content") or "").strip()
    if not text.startswith("###"):
        text = f"### {section}\n{text}"
    cache[key] = {"text": text, "model": resp.get("model", ""),
                  "usage": resp.get("usage", {})}
    save_cache(cache)
    return text


# ═══════════════════════════════════════════════════════════════════════
# 01 招标文件正文 docx
# ═══════════════════════════════════════════════════════════════════════
def _set_cn(style, east_asia: str, ascii_font: str = "Times New Roman",
            size: float | None = None, bold: bool | None = None,
            color=None) -> None:
    from docx.shared import Pt
    style.font.name = ascii_font
    rpr = style.element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", east_asia)
    if size:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color


def _add_para(doc, text: str, size: float = 12, bold: bool = False,
              align=None, indent: bool = True, cn: str = "宋体") -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if not indent:
        p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", cn)
    return p


def _add_markdown_unit(doc, unit_key: str, text: str) -> None:
    """把 LLM/离线生成的 markdown 单元写入 docx（### → Heading 3，列表 → 缩进段落）。"""
    from docx.shared import Pt
    for raw in text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            h = doc.add_heading(line[4:].strip(), level=3)
            h.paragraph_format.first_line_indent = Pt(0)
        elif line.startswith("#### "):
            h = doc.add_heading(line[5:].strip(), level=4)
            h.paragraph_format.first_line_indent = Pt(0)
        elif line.startswith(("- ", "（1）", "1）", "(1)", "1.")):
            p = _add_para(doc, line.strip(), indent=True)
            p.paragraph_format.first_line_indent = Pt(24)
        else:
            _add_para(doc, line.strip(), indent=True)


def _cover(doc) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    for _ in range(4):
        doc.add_paragraph()
    _add_para(doc, "XX市智慧园区建设项目", size=26, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, cn="黑体")
    _add_para(doc, "招 标 文 件", size=22, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, indent=False, cn="黑体")
    doc.add_paragraph()
    for label, value in (
        ("招标编号", TENDER_NO),
        ("招 标 人", "XX市智慧城市管理局"),
        ("招标代理机构", "XX市公共资源交易中心"),
        ("日    期", "2024 年 9 月"),
    ):
        _add_para(doc, f"{label}：{value}", size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    doc.add_page_break()


def _rule_chapter_1(doc) -> None:
    """第一章 招标公告（规则文本）。"""
    doc.add_heading("第一章 招标公告", level=1)
    body = [
        "XX市智慧城市管理局（以下简称'招标人'）就 XX市智慧园区建设项目（招标编号：ZH-2024-018）进行公开招标，欢迎符合条件的投标人参加投标。",
        "一、项目概况：本项目为园区智能化软硬件一体化建设，招标范围包括智慧安防、智慧通行、智慧能耗、智慧停车、一卡通及综合管理平台等系统的供货、安装、调试、培训及质保期服务。",
        "二、投标人资格要求：详见第八章资质要求。",
        "三、招标文件的获取：自本公告发布之日起，投标人可到 XX市公共资源交易中心获取招标文件。",
        "四、投标截止时间：2024 年 10 月 15 日 9 时 30 分；开标时间与地点：同投标截止时间，XX市公共资源交易中心第二开标室。",
        "五、联系方式：招标人联系人：王主任，联系电话：0000-8888888。",
    ]
    for t in body:
        _add_para(doc, t)
    doc.add_page_break()


def _rule_chapter_11(doc) -> None:
    """第十一章 评分标准（规则表格：带边框，find_tables 可稳定识别）。"""
    doc.add_heading("第十一章 评分标准", level=1)
    doc.add_heading("11.1 评标方法", level=2)
    _add_para(doc, "本项目采用综合评分法。评标总分为 100 分，其中技术部分 50 分、商务部分 20 分、价格部分 30 分。评标委员会按照综合得分由高到低排序，推荐中标候选人。")

    doc.add_heading("11.2 技术评分表（50 分）", level=2)
    tech_rows = [
        ("1", "总体技术方案", "10", "方案完整、架构先进合理，完全满足招标要求得 8-10 分；基本合理得 4-7 分；不合理得 0-3 分。"),
        ("2", "平台功能设计", "10", "平台功能覆盖全部功能要求且有人脸识别、设备接入不少于 1000 台等关键能力设计得 8-10 分；覆盖基本功能得 4-7 分；缺项明显得 0-3 分。"),
        ("3", "系统性能", "5", "可用性、并发、响应时间等性能指标完全满足得 4-5 分；基本满足得 2-3 分；不满足得 0-1 分。"),
        ("4", "安全性设计", "5", "等保三级、密码应用、审计等安全设计完整得 4-5 分；基本完整得 2-3 分；缺失得 0-1 分。"),
        ("5", "国产化适配", "5", "提供信创环境部署方案且适配完整得 4-5 分；部分适配得 2-3 分；无方案得 0-1 分。"),
        ("6", "接口与集成", "5", "接口标准符合要求、集成方案完整得 4-5 分；基本符合得 2-3 分；不符合得 0-1 分。"),
        ("7", "项目管理与实施", "5", "实施组织、进度计划科学合理得 4-5 分；基本合理得 2-3 分；不合理得 0-1 分。"),
        ("8", "培训方案", "3", "培训方案完整、学时满足要求得 2-3 分；基本满足得 1 分；缺失得 0 分。"),
        ("9", "售后服务方案", "2", "质保、响应、驻场安排满足要求得 2 分；基本满足得 1 分；缺失得 0 分。"),
    ]
    _add_table(doc, tech_rows)

    doc.add_heading("11.3 商务评分表（20 分）", level=2)
    biz_rows = [
        ("1", "企业资质", "6", "同时具备 ISO9001、ISO27001 认证及 CMMI3 级及以上得 6 分；每缺一项扣 2 分。"),
        ("2", "类似业绩", "8", "近三年类似智慧园区/智慧城市项目每提供 1 个（合同额不低于 500 万元）得 2 分，最高 8 分。"),
        ("3", "项目经理", "4", "项目经理具有 5 年以上智慧园区类项目管理经验且具有 PMP 或信息系统项目管理师证书得 4 分；缺一项扣 2 分。"),
        ("4", "企业信誉", "2", "无失信记录、纳税信用良好的得 2 分；一般得 1 分。"),
    ]
    _add_table(doc, biz_rows)

    doc.add_heading("11.4 价格分计算方法（30 分）", level=2)
    _add_para(doc, "价格分采用低价优先法计算：评标基准价 = 满足招标文件要求且投标价格最低的投标报价。投标报价等于评标基准价的得满分 30 分；其他投标人的价格分统一按照下列公式计算：投标报价得分 =（评标基准价 ÷ 投标报价）× 30。价格分计算保留两位小数。")
    doc.add_page_break()


def _add_table(doc, rows: list[tuple]) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"   # 带边框：pdf_parser find_tables 可稳定识别
    headers = ("序号", "评价项", "分值", "评分细则")
    for i, t in enumerate(headers):
        table.rows[0].cells[i].text = t
    for r in rows:
        cells = table.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = v


def _rule_chapter_12(doc) -> None:
    """第十二章 投标文件格式（规则文本）。"""
    doc.add_heading("第十二章 投标文件格式", level=1)
    doc.add_heading("12.1 装订与份数要求", level=2)
    _add_para(doc, "投标文件应胶装成册，正本 1 份、副本 4 份，另附电子版 1 份（U 盘，内容与正本一致）。正本与副本不一致的，以正本为准。正本封面应有投标人公章及法定代表人签章。")
    doc.add_heading("12.2 目录结构要求", level=2)
    for t in (
        "投标文件应按以下目录结构编制：（1）投标函及投标函附录；（2）法定代表人身份证明及授权委托书；（3）开标一览表；（4）分项报价表；（5）公司介绍；（6）资质证明材料；（7）类似项目业绩证明材料；（8）项目团队人员证明材料；（9）技术方案（含总体架构、功能设计、实施方案、售后服务方案）；（10）其他材料。",
    ):
        _add_para(doc, t)
    doc.add_heading("12.3 格式要求", level=2)
    _add_para(doc, "投标文件统一使用 A4 纸，正文采用宋体小四号字，标题采用黑体；全文连续编排页码；关键证明文件提供原件扫描件并加盖公章。")
    doc.add_page_break()


def build_docx(units_text: dict[str, str], out_path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()
    # 样式：正文宋体小四 1.5 倍行距；标题黑体黑色（默认蓝色需覆盖）
    normal = doc.styles["Normal"]
    _set_cn(normal, "宋体", size=12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Pt(24)
    for name, sz in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13), ("Heading 4", 12)):
        h = doc.styles[name]
        _set_cn(h, "黑体", size=sz, bold=True, color=RGBColor(0, 0, 0))
        h.paragraph_format.first_line_indent = Pt(0)

    _cover(doc)
    _rule_chapter_1(doc)

    # 各章：LLM/离线单元（第十三章在规则章 11/12 之后单独写，保持章节顺序）
    for chapter_title, units in CHAPTER_UNITS:
        if chapter_title.startswith("第十三章"):
            continue
        doc.add_heading(chapter_title, level=1)
        for key, section, _points, _chars in units:
            if key in units_text:
                _add_markdown_unit(doc, key, units_text[key])
        doc.add_page_break()

    _rule_chapter_11(doc)
    _rule_chapter_12(doc)
    for chapter_title, units in CHAPTER_UNITS:
        if chapter_title.startswith("第十三章"):
            doc.add_heading(chapter_title, level=1)
            for key, section, _points, _chars in units:
                if key in units_text:
                    _add_markdown_unit(doc, key, units_text[key])

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"[OK] {out_path.name}")


# ═══════════════════════════════════════════════════════════════════════
# 02 技术规格书 PDF（PyMuPDF 排版）
# ═══════════════════════════════════════════════════════════════════════
SIMSUN = Path("C:/Windows/Fonts/simsun.ttc")
SIMHEI = Path("C:/Windows/Fonts/simhei.ttf")


def _font_ok() -> bool:
    try:
        import fitz
        fitz.Font(fontfile=str(SIMSUN), fontname="F0")
        return True
    except Exception:
        return False


USE_SYSTEM_FONT = _font_ok()


class PdfWriter:
    """PyMuPDF 中文排版小工具：宋体正文 + 黑体标题 + 页码 + 目录书签。"""

    def __init__(self):
        import fitz
        self.doc = fitz.open()   # 1.28 无 new_doc()，open() 创建新文档
        self.page = None
        self.y = 0
        self.page_no = 0
        self.toc: list[list] = []
        self.MARGIN = 72
        self.BOTTOM = 770

    def _new_page(self):
        import fitz
        self.page = self.doc.new_page(width=595, height=842)  # A4
        self.page_no += 1
        self.y = self.MARGIN + 10
        self._footer()

    def _footer(self):
        import fitz
        try:
            self.page.insert_text((self.MARGIN, 810), f"— {self.page_no} —",
                                  fontsize=9, fontname="china-s")
        except Exception:
            pass

    def _text_width(self, text: str, size: float) -> float:
        if USE_SYSTEM_FONT:
            try:
                import fitz
                return fitz.Font(fontfile=str(SIMSUN), fontname="F0").text_length(text, size)
            except Exception:
                pass
        return len(text) * size

    def _draw(self, text: str, size: float, bold: bool, indent: float = 0) -> None:
        import fitz
        x = self.MARGIN + indent
        try:
            fontfile = str(SIMHEI if bold else SIMSUN)
            self.page.insert_text((x, self.y), text, fontsize=size,
                                  fontfile=fontfile, fontname="F0")
        except RuntimeError:
            self.page.insert_text((x, self.y), text, fontsize=size, fontname="china-s")

    def para(self, text: str, size: float = 10.5, bold: bool = False,
             indent: bool = True, spacing: float = 6) -> None:
        """写一段（自动换行 + 分页）。"""
        if self.page is None:
            self._new_page()
        line_w = 595 - 2 * self.MARGIN - (2 * size if indent else 0)
        cur, lines = "", []
        for ch in text:
            if self._text_width(cur + ch, size) > line_w:
                lines.append(cur)
                cur = ch
            else:
                cur += ch
        if cur:
            lines.append(cur)
        for i, ln in enumerate(lines):
            if self.y > self.BOTTOM:
                self._new_page()
            self._draw(ln, size, bold, indent=(2 * size if indent and i == 0 else 0))
            self.y += size + spacing

    def chapter(self, title: str) -> None:
        """新章：黑体标题 + 目录书签。"""
        if self.page is None:
            self._new_page()
        if self.y > self.MARGIN + 60:
            self._new_page()
        self.toc.append([1, title, self.page_no])
        self._draw(title, 16, bold=True)
        self.y += 26

    def param_line(self, name: str, value: str) -> None:
        self.para(f"{name}：{value}", size=10.5, indent=False, spacing=4)

    def finish(self, out_path: Path) -> None:
        if self.toc:
            self.doc.set_toc(self.toc)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(out_path))
        self.doc.close()
        print(f"[OK] {out_path.name}（{self.page_no} 页）")


# 规格书内容：章节 → (段落列表, 参数列表)
# 预埋基线：人脸识别≥99.5% / 存储90天 / MTBF≥50000h / 设备接入≥1000 / 等保三级 /
#           GB/T 28181 / ONVIF / 信创 / UPS≥2h / 车牌≥99%（97% 夜间）/ 采集≤15min
SPEC_CHAPTERS: list[tuple[str, list[str], list[tuple[str, str]]]] = [
    ("第一章 总则", [
        "本技术规格书是 XX市智慧园区建设项目（招标编号：ZH-2024-018）招标文件的重要组成部分，规定了园区智能化系统的主要技术参数、功能要求、接口要求与验收标准。投标人提供的所有设备与软件必须满足本规格书要求，否则视为实质性偏离。",
        "本规格书适用于园区智慧安防、智慧通行、智慧能耗、智慧停车、一卡通及综合管理平台等系统的设计、供货、安装、调试、培训与质保期服务。",
        "投标人应对本规格书逐条响应，并在投标文件技术方案中明确所投产品的品牌、型号、主要技术指标及偏离情况说明。未列明偏离的，视为完全响应。",
        "本规格书未尽事宜，以招标文件正文第四章、第五章要求为准；二者不一致时，以较严格者为准。",
    ], [
        ("项目名称", "XX市智慧园区建设项目"),
        ("招标编号", "ZH-2024-018"),
    ]),
    ("第二章 通用技术要求", [
        "所有设备应适应园区环境条件：工作温度 -10℃～50℃，存储温度 -20℃～70℃，相对湿度 10%～95%（无凝露）。室外设备防护等级不低于 IP66，防雷接地满足 GB 50343 等相关国家标准。",
        "系统平台应支持不少于 1000 台（个）设备的接入管理；平台采用微服务架构，支持横向扩展；单节点故障不影响系统整体运行。",
        "设备平均无故障时间（MTBF）要求：核心服务器、交换机等关键设备 MTBF 不低于 50000 小时；前端设备 MTBF 不低于 30000 小时。",
        "系统应支持信创国产化环境部署，包括但不限于国产 CPU（飞腾、鲲鹏、龙芯等）、国产操作系统（麒麟、统信等）、国产数据库及国产中间件，并提供与信创环境适配的完整证明材料。",
        "所有系统时间应支持 NTP 网络校时，全网设备时间同步误差不超过 1 秒；系统应具备统一运维监控能力，支持对设备在线状态、资源使用率、告警事件的集中监视。",
        "系统整体可用性不低于 99.9%；一般业务操作响应时间不超过 3 秒；复杂统计查询响应时间不超过 10 秒。",
    ], [
        ("设备接入能力", "不少于 1000 台（个）"),
        ("工作温度", "-10℃～50℃"),
        ("室外防护等级", "不低于 IP66"),
        ("系统平均无故障时间 MTBF", "不低于 50000 小时"),
        ("系统可用性", "不低于 99.9%"),
    ]),
    ("第三章 智慧安防系统", [
        "视频监控系统采用全数字网络架构，前端摄像机分辨率不低于 400 万像素，支持 H.265 编码。视频存储采用集中存储方式，录像保存时间不少于 90 天。",
        "摄像机应支持宽动态、背光补偿与红外补光；室外枪机应具备 IP66 防护等级；重点区域摄像机应支持智能分析功能，包括周界入侵检测、区域徘徊检测与物品遗留检测。",
        "人脸识别系统支持白名单比对、陌生人预警与轨迹查询，识别准确率不低于 99.5%；人脸比对响应时间不超过 1 秒。",
        "视频监控接口应符合 GB/T 28181 标准，支持通过国标平台进行统一接入与调阅；录像回放应支持按时间、事件、点位多条件检索。",
        "视频监控系统应支持与门禁、一卡通系统联动，实现报警事件自动弹窗与联动录像；联动响应时间不超过 2 秒。",
    ], [
        ("摄像机分辨率", "不低于 400 万像素"),
        ("视频编码", "H.265"),
        ("录像保存时间", "不少于 90 天"),
        ("人脸识别准确率", "不低于 99.5%"),
        ("人脸比对响应时间", "不超过 1 秒"),
        ("视频监控接口", "GB/T 28181"),
        ("联动响应时间", "不超过 2 秒"),
    ]),
    ("第四章 智慧通行系统", [
        "门禁系统支持刷卡、人脸、二维码等多种认证方式，支持反潜回、防尾随功能；人脸识别终端采用双目活体检测，防止照片、视频攻击。",
        "门禁控制器应支持脱机运行，网络中断时本地存储不少于 10 万条事件记录，网络恢复后自动同步；单台控制器支持不少于 4 个门点。",
        "门禁授权应采用分级管理模式，支持按人员、部门、时间段、门点进行灵活授权；开门响应时间不超过 0.5 秒。",
        "访客管理支持线上预约、访客二维码签发、通行记录查询与黑名单管理；访客通行记录保存不少于 180 天。",
        "通道闸机应具备防夹、防冲撞安全防护功能，断电或紧急情况下自动开闸；闸机通行速度不低于每分钟 35 人。",
    ], [
        ("认证方式", "刷卡 / 人脸 / 二维码"),
        ("活体检测", "双目（可见光 + 红外）"),
        ("单台控制器门数", "不低于 4 门"),
        ("开门响应时间", "不超过 0.5 秒"),
        ("控制器脱机事件存储", "不少于 10 万条"),
        ("访客记录保存", "不少于 180 天"),
    ]),
    ("第五章 智慧能耗系统", [
        "能耗管理系统对园区水、电进行分类分项计量与在线监测，支持能耗趋势分析、异常告警与报表导出；智能电表采集精度不低于 1 级。",
        "数据采集周期不大于 15 分钟；采集终端应支持断点续传，网络中断期间本地缓存数据不少于 7 天。",
        "系统应支持能耗对标分析与定额管理，可按建筑、楼层、区域、部门多维度统计；历史数据保存不少于 3 年。",
        "智能照明系统对公共区域照明实行分时分区控制，支持照度感应、人体感应联动；照明回路应支持远程开关、调光与状态监测。",
        "能耗数据应开放标准接口，支持向园区综合管理平台及上级能耗监管平台推送；接口协议支持 HTTP/HTTPS 与 MQTT。",
    ], [
        ("电表采集精度", "不低于 1 级"),
        ("数据采集周期", "不大于 15 分钟"),
        ("历史数据保存", "不少于 3 年"),
        ("断点续传缓存", "不少于 7 天"),
        ("数据推送协议", "HTTP/HTTPS、MQTT"),
    ]),
    ("第六章 智慧停车系统", [
        "停车管理系统采用车牌识别无感通行，支持车位引导、反向寻车与无人值守收费；车牌识别准确率白天不低于 99%，夜间不低于 97%。",
        "道闸抬杆时间不超过 3 秒；道闸应具备防砸车安全保护功能（地感 + 雷达双重防护）。",
        "车位检测采用地磁与视频联动方式，车位状态上报延迟不超过 10 秒；车位引导屏应实时显示区域空余车位数量。",
        "收费系统支持 ETC、扫码支付、无感支付等多种缴费方式，收费记录保存不少于 3 年；出口无人值守时支持远程呼叫对讲。",
        "停车系统应支持与园区安防平台联动，黑名单车辆自动预警；月租车辆支持线上续费与自动识别放行。",
    ], [
        ("车牌识别准确率", "白天 ≥99%，夜间 ≥97%"),
        ("道闸抬杆时间", "不超过 3 秒"),
        ("车位检测方式", "地磁 + 视频联动"),
        ("车位状态上报延迟", "不超过 10 秒"),
        ("收费记录保存", "不少于 3 年"),
    ]),
    ("第七章 一卡通系统", [
        "一卡通系统实现门禁、消费、考勤、访客一卡通用，支持实体卡与虚拟卡（二维码）两种介质，支持挂失、补卡与消费限额管理。",
        "发卡容量不低于 10 万张；消费交易响应时间不超过 1 秒；脱机状态下消费终端应能正常受理交易。",
        "一卡通数据应采用加密存储与加密传输，卡片密钥实行分级管理；系统应支持与第三方支付平台对接。",
        "考勤管理支持固定班次与弹性班次设置，考勤数据统计应支持按日、周、月自动汇总与报表导出。",
        "一卡通平台应提供统一的管理门户与移动端应用，支持个人消费记录查询、余额查询与挂失自助办理。",
    ], [
        ("卡介质", "实体 IC 卡 + 虚拟卡（二维码）"),
        ("发卡容量", "不低于 10 万张"),
        ("消费交易响应时间", "不超过 1 秒"),
        ("数据安全", "加密存储 + 加密传输"),
    ]),
    ("第八章 网络与机房工程", [
        "园区网络采用核心-汇聚-接入三级架构，核心交换机双机热备；核心交换背板带宽不低于 1.2Tbps，支持万兆上行。",
        "网络应划分为视频专网、办公网、设备网等逻辑区域，区域间通过安全策略隔离；关键链路应支持链路聚合与冗余备份。",
        "机房建设满足 GB 50174 相关要求；UPS 采用在线式，后备时间不少于 2 小时；机房应配备精密空调、环境监测（温湿度、漏水、烟感）与动力环境监控系统。",
        "机房应提供不少于 42U 的标准服务器机柜，机柜供电采用双路供电；弱电间与机房内线缆应按规范做好标识与理线。",
        "室外传输采用单模光纤，光缆芯数按需求 1.5 倍冗余配置；室外管道与手孔井建设满足相关市政规范。",
    ], [
        ("网络架构", "核心-汇聚-接入三级，核心双机热备"),
        ("核心交换背板带宽", "不低于 1.2Tbps"),
        ("UPS 后备时间", "不少于 2 小时"),
        ("机房标准", "GB 50174"),
    ]),
    ("第九章 系统安全要求", [
        "系统安全保护等级不低于网络安全等级保护第三级（等保三级），投标人应在投标文件中提供等保测评配合承诺。",
        "系统应满足商用密码应用安全性评估要求，重要业务数据传输、存储应采用国家密码管理部门认可的密码算法进行加密。",
        "系统应提供完整审计日志，记录用户登录、关键操作、数据访问等行为，审计日志保存不少于 6 个月且不可篡改。",
        "系统应具备入侵检测与防护能力，支持防病毒、防勒索；应定期开展漏洞扫描与渗透测试，及时修复安全漏洞。",
        "数据备份策略：重要业务数据每日全量备份，备份数据异地保存；系统应提供灾难恢复预案，RPO 不超过 24 小时，RTO 不超过 72 小时。",
    ], [
        ("安全等级", "等保三级"),
        ("审计日志保存", "不少于 6 个月"),
        ("数据传输", "全链路加密传输"),
        ("RPO", "不超过 24 小时"),
        ("RTO", "不超过 72 小时"),
    ]),
    ("第十章 接口与验收标准", [
        "视频监控接口应符合 GB/T 28181 标准；设备接入支持 ONVIF 协议；平台对外提供标准 REST API 与消息接口，接口文档应随系统一并交付。",
        "系统应支持与上级政务平台、公安雪亮工程平台等外部系统对接，投标人应免费提供对接开发配合。",
        "验收流程：设备到货验收 → 安装调试完成初验 → 初验合格后进入试运行（3 个月）→ 试运行无重大问题后组织终验。",
        "初验内容应包括：设备数量与型号核对、单点功能测试、系统联调测试、性能指标测试（并发、响应时间、可用性抽测）。",
        "终验依据：招标文件、技术规格书、投标文件及双方确认的设计文件；终验合格后进入质保期（不少于 2 年）。",
        "投标人应提交完整的竣工资料，包括设备清单、点位图、网络拓扑图、系统配置文档、操作手册与维护手册（纸质 3 套 + 电子版 1 套）。",
    ], [
        ("视频监控接口", "GB/T 28181"),
        ("设备接入协议", "ONVIF"),
        ("试运行期", "3 个月"),
        ("质保期", "不少于 2 年"),
        ("竣工资料", "纸质 3 套 + 电子版 1 套"),
    ]),
]


def build_pdf(out_path: Path) -> None:
    w = PdfWriter()
    w.para(f"{TENDER_NAME}技术规格书", size=20, bold=True, indent=False, spacing=16)
    w.para(f"招标编号：{TENDER_NO}", size=12, indent=False, spacing=16)
    for title, paras, params in SPEC_CHAPTERS:
        w.chapter(title)
        for p in paras:
            w.para(p)
        w.para("", spacing=2)
        for name, value in params:
            w.param_line(name, value)
        w.y += 8
    w.finish(out_path)


# ═══════════════════════════════════════════════════════════════════════
# 03 设备清单 xlsx
# ═══════════════════════════════════════════════════════════════════════
HW_ROWS = [
    ("人脸识别终端", "FD-800（双目活体）", "台", 120, 2500, "含安装支架"),
    ("门禁控制器", "AC-2000（四门）", "台", 80, 1800, "含电源箱"),
    ("IC 读卡器", "RD-100", "台", 160, 350, ""),
    ("网络摄像机（枪机）", "IPC-400（400 万）", "台", 260, 900, "室外含护罩"),
    ("网络摄像机（半球）", "IPC-410（400 万）", "台", 180, 850, ""),
    ("硬盘录像机", "NVR-64（64 路）", "台", 12, 6500, ""),
    ("存储服务器", "ST-24（24 盘位）", "台", 3, 85000, "含 8TB 企业级硬盘 ×24"),
    ("应用服务器", "SR-650", "台", 4, 120000, "信创 CPU"),
    ("数据库服务器", "SR-660", "台", 2, 150000, "信创 CPU，双机热备"),
    ("核心交换机", "SW-48X（48 口万兆）", "台", 4, 38000, "双机热备"),
    ("接入交换机", "SW-24（24 口千兆）", "台", 40, 4200, ""),
    ("光模块", "SFP-10G", "个", 96, 550, ""),
    ("能耗采集器", "EM-300", "台", 60, 1200, ""),
    ("智能电表", "MT-100（1 级）", "块", 320, 280, ""),
    ("地磁检测器", "GD-200", "个", 450, 480, ""),
    ("车牌识别一体机", "LPR-900", "套", 24, 9800, "含补光灯"),
    ("道闸", "BG-500（直杆）", "套", 24, 6500, ""),
    ("车位引导屏", "DP-40（双面）", "块", 32, 3200, ""),
    ("发卡器", "CR-100", "台", 10, 1500, ""),
    ("消费机", "POS-200", "台", 30, 2200, ""),
    ("通道闸机", "TG-300（单机芯）", "套", 16, 15000, ""),
    ("UPS 主机", "UPS-60K（60kVA）", "台", 2, 98000, "在线式"),
    ("蓄电池", "BT-100（12V100AH）", "节", 128, 950, ""),
    ("服务器机柜", "CB-42（42U）", "个", 12, 4800, "含 PDU"),
    ("网线", "CAT6 六类", "箱", 120, 750, "305 米/箱"),
    ("室外光纤", "GYTA-12 芯单模", "米", 8000, 4.5, ""),
]
SW_ROWS = [
    ("智慧园区综合管理平台", "V3.2", "套", 1, 480000, "含数据中台、三维可视化"),
    ("视频监控子系统", "V2.5", "套", 1, 120000, ""),
    ("门禁管理子系统", "V2.3", "套", 1, 90000, ""),
    ("能耗管理子系统", "V2.1", "套", 1, 110000, ""),
    ("停车管理子系统", "V2.4", "套", 1, 95000, ""),
    ("一卡通子系统", "V2.2", "套", 1, 85000, ""),
    ("移动端应用", "V1.6（APP/小程序）", "套", 1, 80000, ""),
    ("数据库软件", "国产数据库（信创）", "套", 3, 35000, ""),
    ("中间件", "国产中间件（信创）", "套", 3, 28000, ""),
    ("服务器操作系统", "国产操作系统（信创）", "套", 12, 8000, ""),
    ("人脸识别算法授权", "FD-Algo", "路", 500, 300, "按接入摄像机路数"),
]
SPARE_ROWS = [
    ("读卡器", "RD-100", "台", 20, ""),
    ("摄像机", "IPC-400（400 万）", "台", 16, ""),
    ("光模块", "SFP-10G", "个", 24, ""),
    ("人脸识别终端", "FD-800", "台", 10, ""),
    ("能耗采集器", "EM-300", "台", 8, ""),
    ("地磁检测器", "GD-200", "个", 30, ""),
    ("车牌识别一体机", "LPR-900", "套", 4, ""),
    ("道闸配件", "BG-500P", "套", 4, ""),
    ("智能电表", "MT-100", "块", 24, ""),
    ("蓄电池", "BT-100", "节", 16, ""),
    ("接入交换机", "SW-24", "台", 4, ""),
    ("通道闸机主板", "TG-MB", "块", 4, ""),
    ("发卡器", "CR-100", "台", 3, ""),
    ("消费机", "POS-200", "台", 3, ""),
    ("各类连接线缆", "批", "批", 1, ""),
]


def build_xlsx(out_path: Path) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = Workbook()
    header_font = Font(name="微软雅黑", bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill("solid", fgColor="2F5597")

    def make_sheet(title, headers, rows, widths):
        ws = wb.active if wb.active.title == "Sheet" else wb.create_sheet()
        ws.title = title
        for c, h in enumerate(headers, 1):
            cell = ws.cell(1, c, h)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center")
        for r, row in enumerate(rows, 2):
            for c, v in enumerate(row, 1):
                ws.cell(r, c, v)
        for c, w in enumerate(widths, 1):
            ws.column_dimensions[ws.cell(1, c).column_letter].width = w
        return ws

    make_sheet("硬件清单", ("设备名称", "规格型号", "单位", "数量", "参考单价（元）", "备注"),
               HW_ROWS, (22, 26, 8, 10, 16, 24))
    make_sheet("软件清单", ("软件名称", "版本", "单位", "数量", "参考单价（元）", "备注"),
               SW_ROWS, (26, 24, 8, 10, 16, 30))
    make_sheet("备品备件", ("备件名称", "规格型号", "单位", "数量", "备注"),
               SPARE_ROWS, (22, 26, 8, 10, 24))
    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out_path))
    print(f"[OK] {out_path.name}")


# ═══════════════════════════════════════════════════════════════════════
# 04 补充通知（扫描件）—— PIL 渲染 → 无文本层 PDF（OCR 测试靶）
# ═══════════════════════════════════════════════════════════════════════
def build_scan_pdf(out_path: Path) -> None:
    import fitz
    from PIL import Image, ImageDraw, ImageFont

    font_path = str(SIMHEI)
    pages_content = [
        ("XX市智慧园区建设项目补充通知（一）",
         "各潜在投标人：",
         f"现对 {TENDER_NAME}（招标编号：{TENDER_NO}）招标文件作出如下补充通知：",
         "一、原定投标截止时间 2024 年 10 月 10 日 9 时 30 分，现顺延至 2024 年 10 月 15 日 9 时 30 分，开标时间及地点相应顺延。",
         "二、除上述变更外，招标文件其余内容不变。",
         "三、本补充通知是招标文件的组成部分，与招标文件具有同等法律效力。",
         "特此通知。",
         "招标人：XX市智慧城市管理局",
         "2024 年 9 月 28 日"),
        ("XX市智慧园区建设项目补充通知（二）",
         "各潜在投标人：",
         f"现对 {TENDER_NAME}（招标编号：{TENDER_NO}）招标文件第七章人员要求澄清如下：",
         "一、项目经理须具有 5 年以上智慧园区类项目管理经验，并具有 PMP 或信息系统项目管理师证书。",
         "二、该条款为★条款（实质性要求），不满足将按否决投标处理。",
         "三、投标文件中须提供项目经理劳动合同、社保证明及证书复印件并加盖公章。",
         "特此通知。",
         "招标人：XX市智慧城市管理局",
         "2024 年 10 月 5 日"),
    ]

    doc = fitz.open()
    for title, *lines in pages_content:
        # 渲染"扫描页"位图（A4 @150dpi）
        img = Image.new("RGB", (1240, 1754), "white")
        draw = ImageDraw.Draw(img)
        title_font = ImageFont.truetype(font_path, 40)
        body_font = ImageFont.truetype(font_path, 30)
        small_font = ImageFont.truetype(font_path, 26)
        # 红头标题
        tw = draw.textlength(title, font=title_font)
        draw.text(((1240 - tw) / 2, 90), title, fill=(200, 30, 30), font=title_font)
        draw.line((80, 160, 1160, 160), fill=(200, 30, 30), width=3)
        y = 220
        for ln in lines:
            if ln.startswith("招标人") or ln.startswith("2024"):
                draw.text((90, y), ln, fill=(0, 0, 0), font=small_font)
                y += 60
            else:
                draw.text((90, y), ln, fill=(0, 0, 0), font=body_font)
                y += 70
        # 落款与"公章"示意（模拟扫描件外观）
        draw.text((980, 1420), "（盖章）", fill=(120, 120, 120), font=small_font)
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
            img.save(f.name)
            tmp = Path(f.name)
        try:
            page = doc.new_page(width=595, height=842)
            page.insert_image(page.rect, filename=str(tmp))
        finally:
            tmp.unlink(missing_ok=True)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    doc.close()
    print(f"[OK] {out_path.name}（{len(pages_content)} 页，无文本层）")


# ═══════════════════════════════════════════════════════════════════════
# 样例说明.md
# ═══════════════════════════════════════════════════════════════════════
README_MD = """# 《XX市智慧园区建设项目》样例招标文件包

> 由 `scripts/make_sample_tender.py` 生成。本包是 M1 里程碑（招标文件解析 + 需求提取）的
> 测试靶与验收基线：覆盖 PDF / Word / Excel / 扫描件四类格式，并预埋 15 条量化需求。

## 文件清单

| 文件 | 格式 | 规模 | 测试靶 |
|---|---|---|---|
| 01_招标文件正文.docx | Word | 13 章，约 110 页 | docx 解析 + 需求提取主靶 |
| 02_技术规格书.pdf | PDF（含目录书签） | 10 章，11 页（每章起新页） | 文本 PDF 解析（TOC 标题检测） |
| 03_设备清单.xlsx | Excel | 3 工作表，约 150 行 | xlsx 解析（sheet = page 语义） |
| 04_补充通知(扫描件).pdf | PDF（无文本层） | 2 页 | 扫描页检测（应识别 ocr_pages=[1,2]）+ OCR |

## 预埋需求基线（M1 验收对照）

| # | 类型 | 预埋内容 | 位置 |
|---|---|---|---|
| 1 | 技术要求 | 平台支持人脸识别（人员身份识别） | 正文 4.2 / 规格书 3 |
| 2 | 技术要求 | 设备接入不少于 1000 台（个） | 正文 4.2 / 规格书 2 |
| 3 | 技术要求 | 并发用户数不少于 500 | 正文 4.2 |
| 4 | 技术要求 | 系统可用性不低于 99.9% | 正文 4.4 |
| 5 | 技术要求 | 支持信创国产化环境部署（CPU/OS/数据库/中间件） | 正文 4.4 |
| 6 | 功能要求 | 五大子系统（安防/通行/能耗/停车/一卡通）+ 综合管理平台 | 正文第五章 |
| 7 | 实施要求 | 总工期不超过 12 个月 | 正文 6.1 |
| 8 | 人员要求 | ★项目经理 5 年以上经验 + PMP/信息系统项目管理师 | 正文 7.1 + 扫描件通知二 |
| 9 | 资质要求 | 近三年类似业绩 ≥3 个（单个合同 ≥500 万） | 正文 8.1 |
| 10 | 资质要求 | ★ISO9001 + ISO27001 + CMMI3 级及以上 | 正文 8.1 |
| 11 | 售后服务 | 质保 ≥2 年；响应 ≤2 小时到场；驻场 2 人 | 正文 9.1 |
| 12 | 商务要求 | 总价包干报价；质保金 5% | 正文 10.1 |
| 13 | 评分标准 | 技术 50 / 商务 20 / 价格 30（13 个评分点） | 正文第十一章（规则表格） |
| 14 | 格式要求 | 正本 1 份副本 4 份 + 电子版、胶装成册 | 正文 2.3 / 12.1 |
| 15 | 商务要求 | 投标截止时间顺延至 2024-10-15 | 扫描件通知一 |

## 预埋"不一致"特性（M5 一致性检查的测试靶）

LLM 扩写正文时按真实招标文件样貌自然引入了口径差异，提取层会如实提取并存的两个值：

| 条款 | 正文（docx） | 规格书（PDF） | 说明 |
|---|---|---|---|
| 人脸识别准确率 | 不低于 99%（4.2 扩写句） | 不低于 99.5% | 正文/规格书不一致（真实常见） |
| 质保 | 项目整体质保不少于 2 年 | 质保期不少于 2 年 | 正文另有"硬件原厂质保 3 年、软件免费升级 5 年"并存 |
| 业绩 | 8.1 正式条款"至少完成 3 个" | — | 正文 2.1 另有"近三年至少完成过两项"并存 |

这些并存值**不是 bug**：真实招标文件普遍存在此类口径差异，M5 一致性检查需能识别并提示人工裁决。

## M1 验收标准

1. **解析**：docx 一级章节 ≥ 13；PDF total_pages ≥ 10 且 TOC 标题命中；xlsx 3 个表格块；扫描件 ocr_pages = [1,2]
2. **提取（LLM 模式）**：正文基线 12 条召回 ≥ 10 条；量化数值原样（1000/500/99.9/12/5/2/3/50/20/30/1/4）
3. **提取（离线模式）**：Mock LLM 返回空需求，管线可跑通不报错
4. **评分表规则解析**：13 个评分点（技术 9 + 商务 4），权重和 70；出处标注文件#块号
5. **★条款补扫**：正文 4.2/7.1/8.1 的 ★ 条款 is_star=True 且 importance=高

## 重新生成

```bash
python scripts/make_sample_tender.py            # 完整模式（DeepSeek，可续跑，~10-20 分钟）
python scripts/make_sample_tender.py --no-llm   # 离线模式（内置简版文本）
python scripts/make_sample_tender.py --only scan --no-cache  # 只重建扫描件
```
"""


def build_readme(out_dir: Path) -> None:
    (out_dir / "样例说明.md").write_text(README_MD, encoding="utf-8")
    print("[OK] 样例说明.md")


# ═══════════════════════════════════════════════════════════════════════
# 主流程
# ═══════════════════════════════════════════════════════════════════════
def main() -> None:
    parser = argparse.ArgumentParser(description="样例招标文件包生成器")
    parser.add_argument("--no-llm", action="store_true", help="离线模式（内置简版文本）")
    parser.add_argument("--no-cache", action="store_true", help="忽略缓存重新生成")
    parser.add_argument("--only", default="all",
                        choices=["all", "docx", "pdf", "xlsx", "scan", "readme"],
                        help="只生成指定类别")
    args = parser.parse_args()

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    print(f"输出目录: {OUT_DIR}")

    if args.only in ("all", "docx"):
        use_llm = not args.no_llm
        client = create_llm_client() if use_llm else None
        cache = {} if args.no_cache else load_cache()
        units_text: dict[str, str] = {}
        for chapter, units in CHAPTER_UNITS:
            print(f"[{chapter}]")
            for key, section, points, chars in units:
                units_text[key] = generate_unit(
                    client, key, section, points, chars, use_llm, cache)
        build_docx(units_text, OUT_DIR / "01_招标文件正文.docx")
    if args.only in ("all", "pdf"):
        build_pdf(OUT_DIR / "02_技术规格书.pdf")
    if args.only in ("all", "xlsx"):
        build_xlsx(OUT_DIR / "03_设备清单.xlsx")
    if args.only in ("all", "scan"):
        build_scan_pdf(OUT_DIR / "04_补充通知(扫描件).pdf")
    if args.only in ("all", "readme"):
        build_readme(OUT_DIR)

    print("完成。")


if __name__ == "__main__":
    main()
